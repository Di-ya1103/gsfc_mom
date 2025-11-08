# backend/docgen.py
from docx import Document
from pathlib import Path
from datetime import datetime
import os
import tempfile
import subprocess

def create_docx_and_pdf(uid, mom_sections, metadata, person_summaries, out_dir: Path):
    """
    Produce two files: DOCX and PDF (if possible).
    """
    doc = Document()

    doc.add_heading("Minutes of the Meeting", level=1)
    # metadata table
    doc.add_paragraph(f"Meeting Name: {metadata.get('meeting_name','')}")
    doc.add_paragraph(f"Date of Meeting: {metadata.get('date','')}\t Time: {metadata.get('time','')}")
    doc.add_paragraph(f"Minutes Prepared By: {metadata.get('minutes_prepared_by','')}")
    doc.add_paragraph(f"Location: {metadata.get('location','')}")
    doc.add_paragraph("")

    # 1. Meeting Objective
    doc.add_heading("1. Meeting Objective", level=2)
    doc.add_paragraph(mom_sections.get("meeting_objective",""))

    doc.add_paragraph("")
    doc.add_paragraph("2. Coordinated by: " + metadata.get("coordinated_by",""))

    doc.add_paragraph("")
    doc.add_heading("3. Attendance at Meeting", level=2)
    
    # attendance: metadata 
    attendees = metadata.get("attendees", [])
    if attendees:
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Name"
        hdr_cells[1].text = "Designation/Department"
        hdr_cells[2].text = "E-mail"
        hdr_cells[3].text = "Mobile No."
        for a in attendees:
            row = table.add_row().cells
            row[0].text = a.get("name","")
            row[1].text = a.get("designation","")
            row[2].text = a.get("email","")
            row[3].text = a.get("mobile","")
    else:
        doc.add_paragraph("No attendees provided.")

    doc.add_paragraph("")
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(mom_sections.get("summary",""))

    doc.add_paragraph("")
    doc.add_heading("Action Items / Decisions", level=2)
    for idx, act in enumerate(mom_sections.get("action_items", []), 1):
        doc.add_paragraph(f"{idx}. {act}")

    doc.add_paragraph("")
    doc.add_heading("Detailed Minutes", level=2)
    for seg in mom_sections.get("detailed_minutes", []):
        speaker = seg.get("speaker") or ""
        start = seg.get("start")
        p = doc.add_paragraph()
        p.add_run(f"[{start:.2f}] {speaker} - ").bold = True
        p.add_run(seg.get("text",""))

    # Person wise summary
    doc.add_paragraph("")
    doc.add_heading("Person-wise Summary", level=2)
    for sp, summ in person_summaries.items():
        doc.add_paragraph(f"{sp}: {summ}")

    # Save docx
    out_docx = out_dir / f"{uid}.docx"
    out_pdf = out_dir / f"{uid}.pdf"
    doc.save(out_docx)

    # Try to convert to PDF
    try:
        import docx2pdf
        docx2pdf.convert(str(out_docx), str(out_pdf))
    except Exception as e:
        # fallback: try libreoffice commandline (soffice)
        try:
            subprocess.run(["soffice","--headless","--convert-to","pdf", str(out_docx), "--outdir", str(out_dir)], check=True)
        except Exception:
            # If conversion fails, we just leave pdf missing
            out_pdf = out_docx  # as fallback return docx path
    return out_docx, out_pdf
