import eventlet
eventlet.monkey_patch()  
import os
import json
import base64
import tempfile
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from flask_socketio import SocketIO, emit
from flask_bcrypt import Bcrypt
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from docx import Document
import whisper

#  Initialize Flask App 
app = Flask(__name__)
CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*", async_mode="eventlet")

#  Config 
app.config["JWT_SECRET_KEY"] = os.getenv("JWT_SECRET_KEY", "super-secret-key")  
bcrypt = Bcrypt(app)
jwt = JWTManager(app)

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

users = {}


#  Load Whisper Model 
print("Loading Whisper 'base' model... (10–60 seconds)")
model = whisper.load_model("base")
print("Whisper model loaded successfully!")

@app.route("/api/register", methods=["POST"])
def register():
    try:
        data = request.get_json(force=True)
        username = (data.get("username") or "").strip()
        password = (data.get("password") or "").strip()

        #  Validation
        if not username or not password:
            return jsonify({"error": "All fields are required"}), 400

        #  Check if user already exists
        if username in users:
            return jsonify({"error": "User already exists"}), 400

        #  Hash and save
        hashed_pw = bcrypt.generate_password_hash(password).decode("utf-8")
        users[username] = hashed_pw

        print(f" Registered user: {username}")
        return jsonify({"message": "User registered successfully"}), 201

    except Exception as e:
        print(f"[Register error] {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/login", methods=["POST"])
def login():
    try:
        data = request.get_json(force=True)
        username = (data.get("username") or "").strip()
        password = (data.get("password") or "").strip()

        if username not in users:
            return jsonify({"error": "Invalid username or password"}), 401

        stored_hash = users[username]
        if not bcrypt.check_password_hash(stored_hash, password):
            return jsonify({"error": "Invalid username or password"}), 401

        #  Generate JWT Token
        access_token = create_access_token(identity=username)
        return jsonify({"token": access_token, "username": username}), 200

    except Exception as e:
        print(f"[Login error] {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/profile", methods=["GET"])
@jwt_required()
def profile():
    user = get_jwt_identity()
    return jsonify({"message": f"Welcome, {user['username']}!"})


def create_mom_docx(metadata, transcript, filename):
    
    doc = Document()

    doc.add_heading("Minutes of the Meeting", level=1)
    doc.add_paragraph(f"Meeting Name: {metadata.get('meeting_name', 'N/A')}")
    doc.add_paragraph(f"Date of Meeting: {metadata.get('date', 'N/A')}\t Time: {metadata.get('time', 'N/A')}")
    doc.add_paragraph(f"Minutes Prepared By: {metadata.get('minutes_prepared_by', 'N/A')}")
    doc.add_paragraph(f"Location: {metadata.get('location', 'N/A')}")

    doc.add_paragraph("")  # blank line

    # 1. Meeting Objective
    doc.add_heading("1. Meeting Objective", level=2)
    doc.add_paragraph(metadata.get("meeting_objective", "N/A"))

    # 2. Coordinated by
    doc.add_paragraph("")
    doc.add_heading("2. Coordinated by", level=2)
    doc.add_paragraph(metadata.get("coordinated_by", "N/A"))

    # 3. Attendance
    doc.add_paragraph("")
    doc.add_heading("3. Attendance at Meeting", level=2)
    attendees = ", ".join(metadata.get("attendees", [])) or "No attendees provided."
    doc.add_paragraph(attendees)

    # Summary
    doc.add_paragraph("")
    doc.add_heading("Summary", level=2)
    summary_lines = transcript.strip().split(". ")
    summary_points = summary_lines[:2]  # take first two lines
    for i, s in enumerate(summary_points, 1):
        doc.add_paragraph(f"{i}. {s.strip()}.")

    # Action Items
    doc.add_paragraph("")
    doc.add_heading("Action Items / Decisions", level=2)
    for i, s in enumerate(summary_points, 1):
        doc.add_paragraph(f"{i}. {s.strip()}.")

    # Detailed Minutes
    doc.add_paragraph("")
    doc.add_heading("Detailed Minutes", level=2)
    # Split transcript into pseudo timestamps
    lines = transcript.strip().split(". ")
    for i, line in enumerate(lines):
        doc.add_paragraph(f"[{i*5:.2f}]  - {line.strip()}.")

    # Person-wise Summary 
    doc.add_paragraph("")
    doc.add_heading("Person-wise Summary", level=2)
    doc.add_paragraph("Person1: " + " ".join(lines[:3])[:150] + "...")
    doc.add_paragraph("Person2: " + " ".join(lines[3:6])[:150] + "...")

    # Save file
    path = os.path.join(UPLOAD_DIR, filename)
    doc.save(path)
    return path


@app.route("/api/upload-audio", methods=["POST"])
def upload_audio():

    try:
        if "file" not in request.files:
            return jsonify({"error": "No file uploaded"}), 400

        file = request.files["file"]
        metadata = json.loads(request.form.get("metadata", "{}"))

        filepath = os.path.join(UPLOAD_DIR, file.filename)
        file.save(filepath)
        print(f"🎧 Processing file: {filepath}")

        #  Detect text input 
        ext = os.path.splitext(file.filename)[1].lower()
        transcript = ""

        if ext in [".txt", ".md"]:  
            with open(filepath, "r", encoding="utf-8") as f:
                transcript = f.read().strip()
            print(" Detected text file, skipping Whisper transcription.")
        else:
            # Audio transcription
            print(" Detected audio file, transcribing with Whisper...")
            result = model.transcribe(filepath)
            transcript = result.get("text", "").strip()
            print(" Transcription complete.")

        filename = f"{os.path.splitext(file.filename)[0]}_MoM.docx"
        create_mom_docx(metadata, transcript, filename)

        return jsonify({
            "transcript": transcript,
            "mom_docx": filename,
            "person_summaries": {"Summary": transcript[:500] + "..."}
        })

    except Exception as e:
        print(f"[Upload error] {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/download/<filename>")
def download_file(filename):
 
    try:
        return send_from_directory(UPLOAD_DIR, filename, as_attachment=True)
    except Exception as e:
        print(f"[Download error] {e}")
        return jsonify({"error": str(e)}), 500

 
@socketio.on("audio_chunk")
def handle_audio_chunk(data):
    try:
        blob_data = base64.b64decode(data["blob"])
        with tempfile.NamedTemporaryFile(suffix=".webm", delete=False) as tmp:
            tmp.write(blob_data)
            tmp_path = tmp.name

        result = model.transcribe(tmp_path)
        partial_text = result.get("text", "").strip()
        os.unlink(tmp_path)

        emit("partial_text", {"text": partial_text if partial_text else "[No speech detected]"})

    except Exception as e:
        print(f"[Socket transcription error] {e}")
        emit("partial_text", {"text": f"[Error during transcription: {str(e)}]"})


if __name__ == "__main__":
    print(" Starting GSFC MoM Backend on http://localhost:5001")
    socketio.run(app, host="0.0.0.0", port=5001, debug=True)
